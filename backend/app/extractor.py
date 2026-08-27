"""Multi-format RFQ table/field extractor.

Extracts structured ocean-freight RFQ fields from Excel (.xlsx/.xls), Word
(.docx), PDF (.pdf) and CSV documents. Two complementary strategies run over a
normalised view of the document and their results are merged:

  1. Key-value  - "Label : value" pairs, whether laid out across adjacent cells
     in a form (POL | Shanghai) or as free-text lines ("POL: Shanghai").
  2. Tabular    - a header row followed by data rows (carrier rate matrices);
     the first populated data row is mapped column-by-column.

Header/label text is matched against the canonical synonym dictionary in
``standard_fields.py``, so arbitrary real-world spellings resolve to standard
column names.
"""
import csv
import io
import re
from typing import Dict, List, Optional, Tuple

from .standard_fields import STANDARD_FIELDS, CORE_FIELD_MAP

# Pre-compute (key, synonym, regex) sorted so longer synonyms win first.
_SYNONYMS: List[Tuple[str, str, "re.Pattern"]] = []
for _f in STANDARD_FIELDS:
    for _s in _f["synonyms"]:
        _SYNONYMS.append((_f["key"], _s, re.compile(r"\b" + re.escape(_s) + r"\b", re.IGNORECASE)))
_SYNONYMS.sort(key=lambda t: len(t[1]), reverse=True)


def _clean(text) -> str:
    if text is None:
        return ""
    return re.sub(r"\s+", " ", str(text)).strip()


def _normalise_label(text: str) -> str:
    """Strip trailing colons and parentheticals used in labels."""
    t = _clean(text).rstrip(":").strip()
    return t


# Tokens that are blank-template placeholders, not real values.
_PLACEHOLDERS = {
    "", "x", "n/a", "na", "-", "--", "tbd", "tba", "yes", "no", "yes/no", "yes / no",
    "yes no", "dock", "ground", "0", "$0", "usd", "amount", "rate", "quantity", "basis",
    "unknown", "select", "click here", "insert", "enter", "per container", "per shipment",
    "per day", "% of freight", "% of cif value", "% of value", "business", "residential",
    "jobsite", "true", "false",
}


def _is_placeholder(value: str) -> bool:
    v = _clean(value).lower()
    if v in _PLACEHOLDERS:
        return True
    # bare currency/percent markers or pure punctuation
    if re.fullmatch(r"[$€£%\s.,/\\-]*", v):
        return True
    return False


def _looks_like_label(value: str) -> bool:
    """A cell that is itself a label/section header, not a data value."""
    v = _clean(value)
    if not v:
        return True
    if v.endswith(":"):
        return True
    if _match_field(v) is not None:
        return True
    # ALL-CAPS multi-word section headers with no digits (e.g. "SHIPMENT DETAILS")
    letters = re.sub(r"[^A-Za-z ]", "", v)
    if letters and letters == letters.upper() and " " in letters.strip() and not re.search(r"\d", v):
        return True
    return False


def _is_value(value: str) -> bool:
    return bool(_clean(value)) and not _is_placeholder(value) and not _looks_like_label(value)


def _match_field(label: str) -> Optional[str]:
    """Return the standard field key a label refers to, or None.

    A label qualifies only if it is reasonably short (labels, not sentences) and
    contains a synonym as a whole token. Longest synonym wins.
    """
    label_c = _normalise_label(label)
    if not label_c or len(label_c) > 60:
        return None
    low = label_c.lower()
    for key, syn, rx in _SYNONYMS:
        if rx.search(low):
            # Avoid matching when the synonym is a tiny fragment of a long label.
            if len(syn) >= 2:
                return key
    return None


# ---------------------------------------------------------------------------
# Normalisation: turn any document into (rows, lines)
#   rows  = list of tables, each a list of row-lists of cell strings
#   lines = free-text lines (paragraphs / pdf text) for "Label: value" scanning
# ---------------------------------------------------------------------------
def _from_csv(raw: bytes) -> Tuple[List[List[List[str]]], List[str]]:
    text = raw.decode("utf-8", errors="ignore")
    rows = [[_clean(c) for c in r] for r in csv.reader(io.StringIO(text))]
    return [rows], []


def _from_xlsx(raw: bytes) -> Tuple[List[List[List[str]]], List[str]]:
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    tables = []
    for sn in wb.sheetnames:
        ws = wb[sn]
        rows = [[_clean(c) for c in row] for row in ws.iter_rows(values_only=True)]
        rows = [r for r in rows if any(r)]
        if rows:
            tables.append(rows)
    return tables, []


def _from_docx(raw: bytes) -> Tuple[List[List[List[str]]], List[str]]:
    import docx
    doc = docx.Document(io.BytesIO(raw))
    tables = []
    for t in doc.tables:
        rows = []
        for r in t.rows:
            rows.append([_clean(c.text) for c in r.cells])
        rows = [r for r in rows if any(r)]
        if rows:
            tables.append(rows)
    lines = [_clean(p.text) for p in doc.paragraphs if _clean(p.text)]
    return tables, lines


def _from_pdf(raw: bytes) -> Tuple[List[List[List[str]]], List[str]]:
    import pdfplumber
    tables, lines = [], []
    with pdfplumber.open(io.BytesIO(raw)) as pdf:
        for page in pdf.pages:
            for tb in page.extract_tables() or []:
                rows = [[_clean(c) for c in row] for row in tb]
                rows = [r for r in rows if any(r)]
                if rows:
                    tables.append(rows)
            txt = page.extract_text() or ""
            lines.extend(_clean(l) for l in txt.splitlines() if _clean(l))
    return tables, lines


def _normalise(filename: str, raw: bytes) -> Tuple[str, List[List[List[str]]], List[str]]:
    name = (filename or "").lower()
    if name.endswith(".csv"):
        return "csv", *_from_csv(raw)
    if name.endswith((".xlsx", ".xls")):
        return "xlsx", *_from_xlsx(raw)
    if name.endswith(".docx"):
        return "docx", *_from_docx(raw)
    if name.endswith(".pdf"):
        return "pdf", *_from_pdf(raw)
    raise ValueError(f"Unsupported file type: {filename}")


# ---------------------------------------------------------------------------
# Strategy 1: key-value pairs
# ---------------------------------------------------------------------------
def _kv_from_rows(tables: List[List[List[str]]]) -> Dict[str, str]:
    found: Dict[str, str] = {}
    for rows in tables:
        for ri, row in enumerate(rows):
            for ci, cell in enumerate(row):
                key = _match_field(cell)
                if not key or key in found:
                    continue
                value = ""
                # value to the right in the same row
                for cj in range(ci + 1, len(row)):
                    if _is_value(row[cj]):
                        value = row[cj]
                        break
                    if _match_field(row[cj]):  # ran into the next label; stop
                        break
                # else value directly below in the same column
                if not value and ri + 1 < len(rows):
                    below = rows[ri + 1]
                    if ci < len(below) and _is_value(below[ci]):
                        value = below[ci]
                if value:
                    found[key] = _clean(value)
    return found


def _kv_from_lines(lines: List[str]) -> Dict[str, str]:
    found: Dict[str, str] = {}
    for line in lines:
        m = re.match(r"^\s*([^:]{2,60}?)\s*[:\-]\s*(.+?)\s*$", line)
        if not m:
            continue
        key = _match_field(m.group(1))
        if key and key not in found:
            val = _clean(m.group(2))
            if _is_value(val):
                found[key] = val
    return found


# ---------------------------------------------------------------------------
# Strategy 2: header row + data rows
# ---------------------------------------------------------------------------
def _tabular(tables: List[List[List[str]]]) -> Dict[str, str]:
    best: Dict[str, str] = {}
    best_score = 0
    for rows in tables:
        for hi in range(min(len(rows), 15)):
            header = rows[hi]
            col_map: Dict[int, str] = {}
            for ci, cell in enumerate(header):
                key = _match_field(cell)
                if key and key not in col_map.values():
                    col_map[ci] = key
            if len(col_map) < 3:  # need a real multi-column table, not a form
                continue
            # first data row with content in a mapped column
            for dr in rows[hi + 1:]:
                if not any(dr):
                    continue
                vals = {
                    key: dr[ci] for ci, key in col_map.items()
                    if ci < len(dr) and _is_value(dr[ci])
                }
                if len(vals) > best_score:
                    best_score = len(vals)
                    best = vals
                break
    return best


# ---------------------------------------------------------------------------
# Normalisation of common values (ports, container types, incoterms)
# ---------------------------------------------------------------------------
_PORTS = {
    "shanghai": "Shanghai (CNSHA)", "shenzhen": "Shenzhen (CNSZX)", "ningbo": "Ningbo (CNNGB)",
    "rotterdam": "Rotterdam (NLRTM)", "hamburg": "Hamburg (DEHAM)", "singapore": "Singapore (SGSIN)",
    "los angeles": "Los Angeles (USLAX)", "new york": "New York (USNYC)",
}


def _normalise_values(fields: Dict[str, str]) -> Dict[str, str]:
    for port_key in ("pol", "pod", "place_of_receipt", "place_of_delivery"):
        v = fields.get(port_key)
        if v and "(" not in v:
            low = v.lower()
            for frag, full in _PORTS.items():
                if frag in low or ("lax" in low and port_key == "pod"):
                    fields[port_key] = full
                    break
    if fields.get("container_type"):
        c = fields["container_type"].upper()
        if "40" in c and ("HC" in c or "HIGH" in c):
            fields["container_type"] = "40HC"
        elif "40" in c and ("GP" in c or "STANDARD" in c or "DRY" in c):
            fields["container_type"] = "40GP"
        elif "45" in c:
            fields["container_type"] = "45HC"
        elif "20" in c:
            fields["container_type"] = "20GP"
    if fields.get("incoterms"):
        up = fields["incoterms"].upper()
        match = next((v for v in ("FOB", "CIF", "EXW", "DDP", "FCA", "DAP", "CFR", "CPT") if v in up), None)
        if match:
            fields["incoterms"] = match
    return fields


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------
def extract_document(filename: str, raw: bytes) -> dict:
    """Extract standard RFQ fields from a document.

    Returns:
        {
          "format": "xlsx|docx|pdf|csv",
          "fields": {standard_key: value, ...},   # everything found
          "core":   {pol, pod, container, ...},    # the 8 classic form fields
          "match_count": int,
          "tables": int,
        }
    """
    fmt, tables, lines = _normalise(filename, raw)

    fields: Dict[str, str] = {}
    # Merge order: tabular first (rate matrices), then key-value fills the gaps.
    for src in (_tabular(tables), _kv_from_rows(tables), _kv_from_lines(lines)):
        for k, v in src.items():
            if v and k not in fields:
                fields[k] = v

    fields = _normalise_values(fields)

    core = {form_key: fields.get(std_key, "") for form_key, std_key in CORE_FIELD_MAP.items()}

    return {
        "format": fmt,
        "fields": fields,
        "core": core,
        "match_count": len(fields),
        "tables": len(tables),
    }
